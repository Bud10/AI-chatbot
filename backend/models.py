from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import os
import re
import shutil
from pathlib import Path
import dateparser
from datetime import datetime

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.tools import tool
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import create_retrieval_chain
from langchain_core.documents import Document
from langchain.agents import AgentExecutor, create_tool_calling_agent

from pydantic import BaseModel
from dotenv import load_dotenv
import logging
import PyPDF2
from io import BytesIO
import docx

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # React app URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Set Google API key from environment variable
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in .env file")
os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY

# Initialize LLM and embeddings
llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.2)  # Use gemini-pro (or gemini-1.5-pro if available)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Global variables to store retriever and summary
retriever = None
document_summary = None  # Store the summary of the uploaded document
appointments = []

# Ensure uploads directory exists
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Pydantic model for chat request
class ChatRequest(BaseModel):
    message: str
    session_id: str = "default"  # Added session_id for chat history

# Validation functions
def validate_email(email: str) -> bool:
    return bool(re.match(r"[^@]+@[^@]+\.[^@]+", email))

def validate_phone(phone: str) -> bool:
    return bool(re.match(r"^\+?\d{10,15}$", phone))

# Date parsing function
def parse_date(text: str) -> str:
    text = text.strip().lower().replace("next","")
    text = text.strip()
    parsed = dateparser.parse(text, settings={'PREFER_DATES_FROM': 'future', 'RELATIVE_BASE':datetime.now()})
    if parsed:
        return parsed.strftime("%Y-%m-%d")
    return "Could not parse date."

# Process document (with summarization)
def process_document(file_content: str):
    try:
        # Create document for FAISS
        doc = Document(page_content=file_content)
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = splitter.split_documents([doc])
        if not splits:
            raise ValueError("Document is empty or could not be processed.")
        
        # Create FAISS vector store for retrieval
        vectorstore = FAISS.from_documents(splits, embeddings)
        retriever = vectorstore.as_retriever()
        
        # Generate summary using the LLM
        summary_prompt = f"Summarize the following document in 100-150 words:\n{file_content[:10000]}"  # Limit to avoid token limits
        summary_response = llm.invoke(summary_prompt)
        summary = summary_response.content if hasattr(summary_response, 'content') else str(summary_response)
        
        return {"retriever": retriever, "summary": summary}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing document: {str(e)}")


# Tools
@tool
def query_documents(query: str) -> str:
    """Answer user questions based on the uploaded document."""
    global retriever
    if not retriever:
        return "No document has been uploaded. Please upload a document first."
    rag_prompt = ChatPromptTemplate.from_template("""
    Answer the question based only on the following context:
    <context>
    {context}
    </context>
    Question: {input}
    """)
    rag_chain = create_retrieval_chain(retriever, rag_prompt | llm)
    return rag_chain.invoke({"input": query})["answer"]

@tool
def summarize_document() -> str:
    """Return a summary of the uploaded document."""
    global document_summary
    if not document_summary:
        return "No document has been uploaded or summarized. Please upload a document first."
    return document_summary

@tool
def parse_date_from_text(text: str) -> str:
    """Parse a natural language date (e.g., 'next Monday') into YYYY-MM-DD format."""
    return parse_date(text)

@tool
def schedule_call(name: str, phone: str, email: str) -> str:
    """Schedule a call with the user."""
    if not validate_phone(phone):
        return "Invalid phone number. Please provide a valid phone number (10-15 digits, optional + prefix)."
    if not validate_email(email):
        return "Invalid email address. Please provide a valid email."
    logger.info(f"Scheduled call for {name} at {phone}, {email}")
    return f"Call scheduled successfully for {name} at {phone}, {email}."

@tool
def book_appointment(name: str, email: str, phone: str, date: str, time: str = "") -> str:
    """Book an appointment. Date must be in YYYY-MM-DD format."""
    if not validate_email(email):
        return "Invalid email address. Please provide a valid email."
    if not validate_phone(phone):
        return "Invalid phone number. Please provide a valid phone number (10-15 digits, optional + prefix)."
        
    if date == "Could not parse date.":
        return "Invalid date: could not parse natural language input."
    
    try:
        datetime.strptime(date, "%Y-%m-%d")
    except ValueError:
            return "Invalid date format after parsing. Date must be in YYYY-MM-DD."
        
    
    # Generate unique ID and store appointments
    appointment_id = len(appointments) + 1
    appointment = {
        "id": appointment_id,
        "name": name,
        "email": email,
        "phone": phone,
        "date": date,
        "time": time or "Not specified"
    }
    appointments.append(appointment)
    logger.info(f"Booked appointment : {appointment} ")
    return {f"Appointment booked successfully for {name} on {date} {time}.",
    }

tools = [query_documents, summarize_document, schedule_call, book_appointment, parse_date_from_text]


#Agent setup
system_prompt = """
You are a helpful custom chatbot that can:
- Answer user questions from a user-uploaded document using the query_documents tool.
- Provide a summary of the uploaded document using the summarize_document tool.
- Schedule calls when the user asks to be called, using the schedule_call tool.
- Book appointments when the user requests it, using the book_appointment tool. Use parse_date_from_text to convert natural language dates (e.g., 'next Monday') to YYYY-MM-DD.
- Use parse_date_from_text to convert natural language dates (e.g, 'next monday') to YYYY-MM-DD.
Important guidelines:
- If no document is uploaded, inform the user to upload one before answering document-related questions or summarizing.
- For dates in natural language, always call parse_date_from_text first. 
- If parse_date_from_text fails (returns "Could not parse date.") THEN ask the user for clarification. 
- If parse_date_from_text succeeds, do not ask for additional confirmation unless the user explicitly requests a different year.
- Gather all necessary information (name, phone, email, date, time) through conversation before finalizing a call or appointment.
- If any required parameters for a tool are missing or invalid, ask the user for the correct information before calling the tool again.
- Validate user inputs during the conversation and reprompt if needed.
- Be conversational and helpful.
"""

prompt = ChatPromptTemplate.from_messages([
    ("system", system_prompt),
    MessagesPlaceholder(variable_name="chat_history"),  # Explicitly name the placeholder
    ("human", "{input}"),
    MessagesPlaceholder(variable_name="agent_scratchpad"),  # For intermediate_steps
])

# Chat history setup
message_history = ChatMessageHistory()

# Agent and executor setup
agent = create_tool_calling_agent(llm=llm, tools=tools, prompt=prompt)
agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True, handle_parsing_errors=True)
agent_with_history = RunnableWithMessageHistory(
    runnable=agent_executor,
    get_session_history=lambda session_id: message_history,
    input_messages_key="input",
    history_messages_key="chat_history",
    output_messages_key="output",
)


UPLOAD_DIR = Path(__file__).resolve().parent / "uploads"

# Endpoints
@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    global retriever, document_summary
    try:
        if not file.filename.endswith((".txt", ".pdf", ".docx")):
            raise HTTPException(status_code=400, detail="Only .txt, .pdf, or .docx files are supported.")
        
        # Validate file size (max 10MB)
        if file.size > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 10MB.")
        
        # Read file content
        content_bytes = await file.read()
        
        #Save file to upload folder
        save_path = UPLOAD_DIR / file.filename

        with save_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Extract text based on file type
        try:
            if file.filename.endswith(".txt"):
                content = content_bytes.decode("utf-8")
            elif file.filename.endswith(".pdf"):
                pdf_reader = PyPDF2.PdfReader(BytesIO(content_bytes))
                content = "".join(page.extract_text() or "" for page in pdf_reader.pages)
            elif file.filename.endswith(".docx"):
                doc = docx.Document(BytesIO(content_bytes))
                content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
        
        # Process document and generate summary
        result = process_document(content)
        retriever = result["retriever"]
        document_summary = result["summary"]
        
        return JSONResponse(content={
            "message": "Document uploaded, processed, and summarized successfully",
            "summary": document_summary
        })
    except Exception as e:
        logger.error(f"Error in upload_document: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/chat")
async def chat(request: ChatRequest):
    try:
        response = agent_with_history.invoke(
            {"input": request.message, "chat_history": []},
            config={"configurable": {"session_id": request.session_id}}
        )
        output = response.get("output", "Sorry, something went wrong. Please try again.")
        if not output:
            logger.error("Empty output from agent_executor")
            output = "Sorry, I couldn't process your request. Please try again."
        
        # response_data = {"response": output}
        # if "Appointment booked successfully" in output and appointments:
        #     response_data["appointments"] = appointments

        return JSONResponse(content={"response": response["output"]})
    except Exception as e:
        logger.error(f"Error in chat: {str(e)}")
        import traceback
        traceback.print_exc()
        return JSONResponse(content={"response": f"Error: {str(e)}"}, status_code=500)
    
@app.get("/appointments")
async def get_appointments():
    try:
        return JSONResponse(content={"appointments": appointments})
    except Exception as e:
        logger.error(f"Error in get_appointments: {str(e)}")
        return JSONResponse(content={"error": str(e)}, status_code=500)
    



